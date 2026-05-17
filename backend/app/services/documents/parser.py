import os
import io
from typing import Optional
from pypdf import PdfReader
from docx import Document
from backend.app.core.logging import logger
from backend.app.core.config import settings
from backend.app.services.documents.ocr import ocr_service

class DocumentParser:
    """
    Unified textual extraction system capable of ingesting and sanitizing 
    heterogeneous file byte-streams. Supports high-fidelity Vision OCR.
    """
    
    @staticmethod
    async def parse_file(file_bytes: bytes, filename: str) -> str:
        """
        Inspects file extensions and dispatches appropriate extraction logic.
        
        :param file_bytes: Unprocessed raw binary content of the file.
        :param filename: The source filename, used to infer file format mappings.
        :returns: Clean, formatted textual representation.
        """
        ext = os.path.splitext(filename.lower())[1]
        logger.info(f"[Doc-Parser] Ingesting {filename} | Extension: {ext} | Size: {len(file_bytes)} bytes.")
        
        try:
            if ext == '.pdf':
                return await DocumentParser._parse_pdf(file_bytes)
            elif ext in ['.docx', '.doc']:
                return DocumentParser._parse_docx(file_bytes)
            elif ext in ['.txt', '.md', '.markdown', '.json', '.py', '.js', '.ts', '.csv']:
                return DocumentParser._parse_plaintext(file_bytes)
            else:
                # Default fallback attempt to parse as plaintext
                logger.warning(f"[Doc-Parser] Unknown extension '{ext}'. Defaulting to Plaintext extraction fallback.")
                return DocumentParser._parse_plaintext(file_bytes)
        except Exception as e:
            logger.error(f"[Doc-Parser] FAILED parsing document '{filename}': {e}", exc_info=True)
            raise ValueError(f"Unsupported or corrupted file format ({filename}): {str(e)}")

    @staticmethod
    async def _parse_pdf(file_bytes: bytes) -> str:
        """Extracts textual data from PDF, prioritizing Vision OCR for high fidelity."""
        
        # Step 1: Attempt Deep Vision OCR if enabled (Key lookup happens inside ocr_service)
        if settings.USE_VISION_OCR:
            logger.info("[Doc-Parser] Initiating Deep Vision OCR pipeline via GPT-4o-mini.")
            ocr_text = await ocr_service.process_pdf(file_bytes)
            if ocr_text:
                logger.info("[Doc-Parser] Deep Vision OCR successful.")
                return ocr_text
            logger.warning("[Doc-Parser] Deep Vision OCR returned empty or failed. Falling back to pypdf.")

        # Step 2: Fallback to standard unicode extraction
        stream = io.BytesIO(file_bytes)
        reader = PdfReader(stream)
        
        extracted_pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            text_clean = text.strip()
            if text_clean:
                extracted_pages.append(f"--- PAGE {i+1} ---\n{text_clean}")
                
        if not extracted_pages:
            raise ValueError("Zero extractable textual characters encountered. Document might be image-only/scanned and Vision OCR is disabled/failed.")
            
        return "\n\n".join(extracted_pages)

    @staticmethod
    def _parse_docx(file_bytes: bytes) -> str:
        """Extracts paragraphs and basic table formats using python-docx."""
        stream = io.BytesIO(file_bytes)
        doc = Document(stream)
        
        extracted_elements = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                extracted_elements.append(text)
                
        # Support lightweight extraction from tables present in the word document
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    extracted_elements.append(row_text)
                    
        return "\n\n".join(extracted_elements)

    @staticmethod
    def _parse_plaintext(file_bytes: bytes) -> str:
        """Decodes standard byte blocks into UTF-8, falling back to latin-1 upon failure."""
        try:
            return file_bytes.decode("utf-8", errors="replace")
        except Exception:
            # Highly resilient decoding fallback for older document files
            return file_bytes.decode("latin-1", errors="replace")

# Functional Export
document_parser = DocumentParser()
